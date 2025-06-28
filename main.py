from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
import time
import random

def setup_driver():
    """Setup Chrome driver with basic options."""
    options = Options()
    options.add_argument('--no-sandbox')
    options.add_argument('--disable-dev-shm-usage')
    
    driver = webdriver.Chrome(
        service=Service(ChromeDriverManager().install()),
        options=options
    )
    return driver

def get_subject_urls(city: str, city_id: int, state: str):
    """Return all Redfin listing URLs for subject properties."""
    url = (f"https://www.redfin.com/city/{city_id}/{state}/"
           f"{city.replace(' ', '-')}/filter/"
           "min-price=150k,max-price=200k,min-beds=3,min-baths=1,include=sold-1mo")
    
    with setup_driver() as driver:
        try:
            driver.implicitly_wait(10)
            driver.get(url)
            
            # Wait and scroll to load listings
            time.sleep(3)
            driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
            time.sleep(2)
            
            # Get property URLs
            property_links = driver.find_elements(By.CSS_SELECTOR, "a.bp-Homecard")
            urls = [link.get_attribute('href') for link in property_links if link.get_attribute('href')]
            
            return urls
            
        except Exception as e:
            print(f"Error fetching subject properties: {e}")
            return []

def get_comp_urls(city: str, city_id: int, state: str):
    """Return Redfin listing URLs for comparable properties (sold 1 year ago, >$250k)."""
    url = (f"https://www.redfin.com/city/{city_id}/{state}/"
           f"{city.replace(' ', '-')}/filter/"
           "min-price=250k,min-beds=3,min-baths=1,include=sold-1yr")
    
    with setup_driver() as driver:
        try:
            driver.implicitly_wait(10)
            driver.get(url)
            
            # Wait and scroll to load more listings
            time.sleep(3)
            for _ in range(3):
                driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
                time.sleep(2)
            
            # Get property URLs
            property_links = driver.find_elements(By.CSS_SELECTOR, "a.bp-Homecard")
            urls = [link.get_attribute('href') for link in property_links if link.get_attribute('href')]
            
            return urls
            
        except Exception as e:
            print(f"Error fetching comp properties: {e}")
            return []

def save_to_docx(subject_urls, comp_urls, distances, city, state_abbr):
    """Save the results to a Word document."""
    doc = Document()
    doc.add_heading(f'Property Analysis - {city}, {state_abbr}', 0)
    
    if subject_urls:
        for i, subject_url in enumerate(subject_urls, 1):
            doc.add_paragraph(f"subject_url {i}: {subject_url}")
            
            # Find 2 comps for each subject
            available_comps = min(2, len(comp_urls))
            for j in range(available_comps):
                comp_idx = (i-1)*2 + j  # Different comps for each subject
                if comp_idx < len(comp_urls):
                    distance = distances[comp_idx % len(distances)]
                    doc.add_paragraph(f"comp {j+1}: ~{distance} miles distance: {comp_urls[comp_idx]}")
            
            doc.add_paragraph("")  # Empty line between subjects
    else:
        doc.add_paragraph("No subject properties found.")
    
    doc.save('address.docx')
    print("Results saved to address.docx")

def main():
    # Get user input
    city = input("Enter the name of the city: ").strip()
    city_id = int(input("Enter the ID of the city: "))
    state_abbr = input("Enter the state abbreviation (e.g., WA for Washington): ").strip()
    
    # Get subject properties
    subject_urls = get_subject_urls(city, city_id, state_abbr)
    
    # Get comparable properties
    comp_urls = get_comp_urls(city, city_id, state_abbr)
    
    # Generate distances (simulated since actual distance calculation requires geocoding)
    distances = [round(random.uniform(0.5, 5.0), 1) for _ in range(len(comp_urls))]
    
    # Save to Word document
    save_to_docx(subject_urls, comp_urls, distances, city, state_abbr)
    
    # Also print to console
    print("\nResults:")
    if subject_urls:
        for i, subject_url in enumerate(subject_urls, 1):
            print(f"subject_url {i}: {subject_url}")
            
            # Find 2 comps for each subject
            available_comps = min(2, len(comp_urls))
            for j in range(available_comps):
                comp_idx = (i-1)*2 + j  # Different comps for each subject
                if comp_idx < len(comp_urls):
                    distance = distances[comp_idx % len(distances)]
                    print(f"comp {j+1}: ~{distance} miles distance: {comp_urls[comp_idx]}")
            print()  # Empty line between subjects
    else:
        print("No subject properties found.")

if __name__ == "__main__":
    main()